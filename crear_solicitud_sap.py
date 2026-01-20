"""
Script para generar documento Word de solicitud SAP
"""
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime

# Crear documento
doc = Document()

# Título
titulo = doc.add_heading('Solicitud de Acceso RFC para Integración Dashboard Logístico', 0)
titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Fecha
fecha = doc.add_paragraph(f'Fecha: {datetime.now().strftime("%d de enero de %Y")}')
fecha.alignment = WD_ALIGN_PARAGRAPH.RIGHT

doc.add_paragraph()

# Destinatario
doc.add_paragraph('Estimado equipo de TI,')
doc.add_paragraph()
doc.add_paragraph('Estamos desarrollando un dashboard de inteligencia logística para Import Aceros S.A. que requiere conexión de SOLO LECTURA a SAP. A continuación detallamos los requerimientos técnicos:')

# Sección 1: Credenciales
doc.add_heading('1. CREDENCIALES DE CONEXIÓN REQUERIDAS', level=1)
doc.add_paragraph('Por favor proporcionar los siguientes datos de conexión:')

# Tabla de credenciales
table1 = doc.add_table(rows=6, cols=3)
table1.style = 'Table Grid'

# Encabezados
hdr_cells = table1.rows[0].cells
hdr_cells[0].text = 'Campo'
hdr_cells[1].text = 'Descripción'
hdr_cells[2].text = 'Valor (completar)'

# Datos
datos_cred = [
    ('SAP_HOST', 'IP o hostname del servidor SAP', ''),
    ('SAP_SYSNR', 'Número de sistema', ''),
    ('SAP_CLIENT', 'Mandante (ej: 100, 300)', ''),
    ('SAP_USER', 'Usuario para conexión RFC', ''),
    ('SAP_PASSWORD', 'Contraseña del usuario', ''),
]

for i, (campo, desc, valor) in enumerate(datos_cred, 1):
    row = table1.rows[i].cells
    row[0].text = campo
    row[1].text = desc
    row[2].text = valor

doc.add_paragraph()

# Sección 2: Usuario RFC
doc.add_heading('2. USUARIO RFC REQUERIDO', level=1)
doc.add_paragraph('Crear un usuario de tipo Sistema/RFC con permisos de SOLO LECTURA a las siguientes tablas:')

# Tabla de permisos
table2 = doc.add_table(rows=8, cols=3)
table2.style = 'Table Grid'

hdr_cells2 = table2.rows[0].cells
hdr_cells2[0].text = 'Tabla SAP'
hdr_cells2[1].text = 'Descripción'
hdr_cells2[2].text = 'Uso en Dashboard'

datos_tablas = [
    ('MARD', 'Stock por almacén', 'Inventario actual'),
    ('MAKT', 'Textos de material', 'Descripciones de productos'),
    ('EKKO', 'Cabecera pedidos compra', 'Órdenes abiertas'),
    ('EKPO', 'Posiciones pedidos compra', 'Detalle de órdenes'),
    ('LFA1', 'Maestro proveedores', 'Datos de proveedores'),
    ('VBAK', 'Cabecera pedidos venta', 'Ventas históricas'),
    ('VBAP', 'Posiciones pedidos venta', 'Detalle de ventas'),
]

for i, (tabla, desc, uso) in enumerate(datos_tablas, 1):
    row = table2.rows[i].cells
    row[0].text = tabla
    row[1].text = desc
    row[2].text = uso

doc.add_paragraph()

# Sección 3: Preguntas técnicas
doc.add_heading('3. PREGUNTAS TÉCNICAS', level=1)

preguntas = [
    '¿SAP es on-premise o S/4HANA Cloud?',
    '¿Está habilitado el módulo RFC para conexiones externas?',
    '¿Se requiere conexión VPN o se puede acceder por IP pública?',
    '¿Tienen SAP NetWeaver RFC SDK disponible para descargar?',
    '¿Cuál es el puerto de conexión RFC? (por defecto 3300)',
]

for pregunta in preguntas:
    p = doc.add_paragraph(pregunta, style='List Bullet')

doc.add_paragraph()

# Sección 4: Alternativa Cloud
doc.add_heading('4. ALTERNATIVA (Si es SAP S/4HANA Cloud)', level=1)
doc.add_paragraph('Si utilizan SAP S/4HANA Cloud, en lugar de RFC necesitamos:')

alternativas = [
    'Acceso a APIs OData del Business Hub',
    'Communication Arrangement configurado',
    'Client ID y Client Secret para autenticación OAuth',
    'URL del API Gateway',
]

for alt in alternativas:
    doc.add_paragraph(alt, style='List Bullet')

doc.add_paragraph()

# Nota importante
doc.add_heading('NOTA DE SEGURIDAD', level=1)
nota = doc.add_paragraph()
nota.add_run('El dashboard realizará únicamente consultas de lectura. ').bold = True
nota.add_run('Nunca se modificarán, eliminarán o crearán datos en SAP. El acceso es exclusivamente para visualización de información en tiempo real.')

doc.add_paragraph()
doc.add_paragraph()

# Cierre
doc.add_paragraph('Quedamos atentos a su respuesta.')
doc.add_paragraph()
doc.add_paragraph('Atentamente,')
doc.add_paragraph()
doc.add_paragraph('_____________________________')
doc.add_paragraph('[Nombre del solicitante]')
doc.add_paragraph('Proyecto: Dashboard Cerebro de Acero')
doc.add_paragraph('Import Aceros S.A.')

# Guardar documento
output_path = r'c:\Users\LENOVO\Documents\ROBOT IMPORT ACEROS\Demo_Import_Aceros\SOLICITUD_ACCESO_SAP.docx'
doc.save(output_path)
print(f"✅ Documento creado: {output_path}")
